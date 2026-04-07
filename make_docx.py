"""Dissertatsiya.md ni Word (.docx) formatiga o'girish."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

MD_PATH = Path(__file__).parent / "dissertatsiya.md"
DOCX_PATH = Path(__file__).parent / "dissertatsiya.docx"

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

# ── Styles ──
style_normal = doc.styles["Normal"]
style_normal.font.name = "Times New Roman"
style_normal.font.size = Pt(14)
style_normal.paragraph_format.line_spacing = 1.5
style_normal.paragraph_format.first_line_indent = Cm(1.25)
style_normal.paragraph_format.space_after = Pt(0)
style_normal.paragraph_format.space_before = Pt(0)

# Headings
for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Times New Roman"
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)
    hs.paragraph_format.first_line_indent = Cm(0)
    if level == 1:
        hs.font.size = Pt(16)
        hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        hs.font.size = Pt(15)
    else:
        hs.font.size = Pt(14)


def add_table(headers, rows):
    """Jadval qo'shish."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    doc.add_paragraph()  # space after table


def parse_table_block(lines):
    """Markdown jadval satrlarini parse qilish."""
    headers = []
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if all(set(c) <= set("- :") for c in cells):
            continue  # separator line
        if not headers:
            headers = cells
        else:
            rows.append(cells)
    return headers, rows


def add_paragraph_with_bold(text):
    """**bold** qismlarni to'g'ri formatlash."""
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p


# ── Parse markdown ──
md_text = MD_PATH.read_text(encoding="utf-8")
lines = md_text.split("\n")

i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Skip empty lines and horizontal rules
    if not stripped or stripped == "---":
        i += 1
        continue

    # Headings
    if stripped.startswith("# ") and not stripped.startswith("## "):
        # Main title
        title = stripped[2:].strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "Times New Roman"
        doc.add_paragraph()  # empty line after title
        i += 1
        continue

    if stripped.startswith("## "):
        text = stripped[3:].strip()
        doc.add_heading(text, level=1)
        i += 1
        continue

    if stripped.startswith("### "):
        text = stripped[4:].strip()
        doc.add_heading(text, level=2)
        i += 1
        continue

    # Code blocks — format as indented text
    if stripped.startswith("```"):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith("```"):
            code_lines.append(lines[i])
            i += 1
        i += 1  # skip closing ```

        for cl in code_lines:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(cl)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        doc.add_paragraph()
        continue

    # Tables
    if stripped.startswith("|"):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            table_lines.append(lines[i])
            i += 1
        headers, rows = parse_table_block(table_lines)
        if headers and rows:
            add_table(headers, rows)
        continue

    # Numbered lists
    if re.match(r"^\d+\.\s", stripped):
        text = re.sub(r"^\d+\.\s", "", stripped)
        p = doc.add_paragraph(style="List Number")
        # Handle bold
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        i += 1
        continue

    # Bullet lists
    if stripped.startswith("- "):
        text = stripped[2:]
        p = doc.add_paragraph(style="List Bullet")
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        i += 1
        continue

    # Regular paragraph (with bold support)
    if stripped:
        add_paragraph_with_bold(stripped)

    i += 1

# ── Save ──
doc.save(str(DOCX_PATH))
print(f"Tayyor: {DOCX_PATH}")
